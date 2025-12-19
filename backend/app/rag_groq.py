import json
import os
from io import BytesIO
from typing import List, Tuple, Optional

import numpy as np
import httpx
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from sqlmodel import select
from fastapi import HTTPException, UploadFile, status
from sentence_transformers import SentenceTransformer
from dotenv import load_dotenv

from .db import get_session
from .models import Document, DocChunk

# -------------------------------------------------------------------
# Env + Groq config
# -------------------------------------------------------------------
load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
GROQ_BASE_URL = "https://api.groq.com/openai/v1"


def _check_key():
    """Check if the GROQ API key is properly configured."""
    if not GROQ_API_KEY or not GROQ_API_KEY.strip():
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=(
                "GROQ API key not properly configured. "
                "Please set the GROQ_API_KEY environment variable in your .env file. "
                "You can get a free API key from https://console.groq.com/keys"
            ),
        )


# -------------------------------------------------------------------
# Embedding model (FAST MODEL)
# -------------------------------------------------------------------
_embedding_model: Optional[SentenceTransformer] = None
_EMBEDDING_MODEL_NAME = "sentence-transformers/paraphrase-MiniLM-L3-v2"


def _get_embedding_model() -> SentenceTransformer:
    global _embedding_model
    if _embedding_model is None:
        try:
            _embedding_model = SentenceTransformer(_EMBEDDING_MODEL_NAME)
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=(
                    f"Failed to load embedding model '{_EMBEDDING_MODEL_NAME}': {str(e)}. "
                    "Make sure sentence-transformers is installed and you have internet access for the first run."
                ),
            )
    return _embedding_model


def embed_texts(texts: List[str]) -> List[List[float]]:
    if not texts:
        return []
    try:
        model = _get_embedding_model()
        embeddings = model.encode(texts, convert_to_numpy=True, show_progress_bar=False)
        if len(embeddings.shape) == 1:
            return [embeddings.tolist()]
        return [emb.tolist() for emb in embeddings]
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate embeddings: {str(e)}",
        )


# -------------------------------------------------------------------
# GROQ Chat
# -------------------------------------------------------------------
def chat_answer(prompt: str, max_retries: int = 3) -> str:
    """Send a prompt to the GROQ API with retry logic and better error handling."""
    _check_key()

    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json",
    }

    payload = {
        "model": "llama-3.1-8b-instant",
        "messages": [
            {
                "role": "system",
                "content": "You are Foresight AI, a helpful knowledge assistant.",
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.2,
    }

    last_error = None
    for attempt in range(max_retries):
        try:
            with httpx.Client(timeout=60.0) as client:
                resp = client.post(
                    f"{GROQ_BASE_URL}/chat/completions",
                    headers=headers,
                    json=payload,
                    timeout=60.0,
                )

                if resp.status_code == 200:
                    data = resp.json()
                    if "choices" in data and data["choices"]:
                        return data["choices"][0]["message"]["content"]
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail="Invalid response format from GROQ API",
                    )

                if resp.status_code == 429 or resp.status_code >= 500:
                    retry_after = int(resp.headers.get("Retry-After", 5)) + 1
                    import time
                    time.sleep(retry_after)
                    continue

                error_detail = resp.text
                try:
                    error_data = resp.json()
                    error_detail = error_data.get("error", {}).get("message", str(error_data))
                except Exception:
                    pass

                raise HTTPException(
                    status_code=resp.status_code,
                    detail=f"GROQ API error: {error_detail}",
                )

        except httpx.RequestError as e:
            last_error = f"Request failed: {str(e)}"
            if attempt == max_retries - 1:
                raise HTTPException(
                    status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
                    detail=f"Failed to connect to GROQ API after {max_retries} attempts: {last_error}",
                )
            continue

    raise HTTPException(
        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
        detail=f"Failed to get response from GROQ API after {max_retries} attempts: {last_error}",
    )


# -------------------------------------------------------------------
# Text / chunks / DB helpers
# -------------------------------------------------------------------
def chunk_text(text: str, max_chars: int = 1200, overlap: int = 250) -> List[str]:
    text = text.replace("\r", " ")
    chunks: List[str] = []
    start = 0
    n = len(text)
    while start < n:
        end = min(start + max_chars, n)
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == n:
            break
        start = max(0, end - overlap)
    return chunks


def extract_text_from_upload(upload: UploadFile) -> str:
    """Extract text from uploaded file with improved error handling."""
    if not upload or not upload.filename:
        raise ValueError("No file provided or invalid file")

    filename = upload.filename.lower()
    try:
        data = upload.file.read()
        if not data:
            raise ValueError("Empty file provided")

        if filename.endswith(".pdf"):
            try:
                doc = fitz.open(stream=data, filetype="pdf")
                parts = []
                max_pages = min(len(doc), 80)
                for i in range(max_pages):
                    page = doc.load_page(i)
                    text = page.get_text()
                    if text.strip():
                        parts.append(text)
                return "\n".join(parts) if parts else ""
            except Exception as e:
                print(f"Error processing PDF: {str(e)}")
                return ""

        elif filename.endswith((".docx", ".doc")):
            try:
                bio = BytesIO(data)
                d = DocxDocument(bio)
                text = "\n".join(p.text for p in d.paragraphs if p.text.strip())
                return text if text.strip() else ""
            except Exception as e:
                print(f"Error processing Word document: {str(e)}")
                return ""

        else:
            try:
                text = data.decode("utf-8", errors="replace")
                return text if text.strip() else ""
            except Exception as e:
                print(f"Error processing text file: {str(e)}")
                return ""

    except Exception as e:
        print(f"Error reading file: {str(e)}")
        return ""


def add_document_and_chunks(name: str, text: str) -> int:
    """Add document and its chunks to the database with improved error handling."""
    if not text or not text.strip():
        print("Error: No text content to process")
        return 0

    try:
        chunks = chunk_text(text)
        if not chunks:
            print("Error: No valid chunks could be created from the text")
            return 0

        vectors = embed_texts(chunks)

        from .models import Document, DocChunk  # type: ignore

        with get_session() as session:
            doc = Document(name=name)
            session.add(doc)
            session.commit()
            session.refresh(doc)

            for i, (content, vec) in enumerate(zip(chunks, vectors)):
                try:
                    session.add(
                        DocChunk(
                            document_id=doc.id,
                            content=content,
                            embedding=json.dumps(vec),
                        )
                    )
                except Exception as e:
                    print(f"Error adding chunk {i+1}: {str(e)}")
                    continue

            session.commit()
            return len(chunks)

    except Exception as e:
        print(f"Error in add_document_and_chunks: {str(e)}")
        from .models import Document  # type: ignore
        with get_session() as session:
            doc = session.exec(select(Document).where(Document.name == name)).first()
            if doc:
                session.delete(doc)
                session.commit()
        return 0


def list_documents():
    from .models import Document, DocChunk  # type: ignore
    with get_session() as session:
        docs = session.exec(select(Document)).all()
        result = []
        for d in docs:
            chunks = session.exec(select(DocChunk).where(DocChunk.document_id == d.id)).all()
            count = len(chunks)
            result.append({"id": d.id, "name": d.name, "chunks": count})
        return result


def delete_document(doc_id: int):
    from .models import Document, DocChunk  # type: ignore
    with get_session() as session:
        doc = session.get(Document, doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail="Document not found")
        chunks = session.exec(select(DocChunk).where(DocChunk.document_id == doc_id)).all()
        for c in chunks:
            session.delete(c)
        session.delete(doc)
        session.commit()


def _cosine(a: np.ndarray, b: np.ndarray) -> float:
    denom = float(np.linalg.norm(a) * np.linalg.norm(b))
    if denom == 0:
        return 0.0
    return float(np.dot(a, b) / denom)


def retrieve_relevant_chunks(question: str, top_k: int = 5):
    from .models import DocChunk  # type: ignore
    q_vec_list = embed_texts([question])
    if not q_vec_list:
        return []
    q_vec = np.array(q_vec_list[0], dtype=float)
    with get_session() as session:
        all_chunks = session.exec(select(DocChunk)).all()
    scored = []
    for ch in all_chunks:
        vec = np.array(json.loads(ch.embedding), dtype=float)
        score = _cosine(q_vec, vec)
        scored.append((ch, score))
    scored.sort(key=lambda x: x[1], reverse=True)
    return scored[:top_k]


def build_context_with_citations(chunks_with_scores) -> str:
    lines = []
    for ch, score in chunks_with_scores:
        lines.append(f"[doc_id={ch.document_id} score={score:.3f}] {ch.content}")
    return "\n\n".join(lines)


def summarize_document(doc_id: int) -> str:
    from .models import DocChunk  # type: ignore
    with get_session() as session:
        chunks = session.exec(select(DocChunk).where(DocChunk.document_id == doc_id)).all()
    if not chunks:
        raise HTTPException(status_code=404, detail="No chunks for this document")
    text = "\n".join(c.content for c in chunks)
    prompt = f"Summarize the following internal document in 6-8 formal bullet points:\n{text}"
    return chat_answer(prompt)


def generate_quiz_for_document(doc_id: int) -> str:
    from .models import DocChunk  # type: ignore
    with get_session() as session:
        chunks = session.exec(select(DocChunk).where(DocChunk.document_id == doc_id)).all()
    if not chunks:
        raise HTTPException(status_code=404, detail="No chunks for this document")
    text = "\n".join(c.content for c in chunks[:15])
    prompt = (
        "Based on the following document content, generate 5 MCQ-style quiz questions "
        "with 4 options each and clearly mark the correct answer:\n"
        f"{text}"
    )
    return chat_answer(prompt)
